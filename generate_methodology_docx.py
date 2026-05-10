"""Generate methodology Word document for the RGB Crop Maturity ML manuscript."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"C:\CaudeProjects\MaturityML\Methodology_Section.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)

def set_font(run, bold=False, italic=False, size=12):
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=14)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=12)
    return p

def body(text, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def body_mixed(parts, justify=True):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("2. Materials and Methods")
set_font(r, bold=True, size=16)

# ══════════════════════════════════════════════════════════════════════════════
# 2.1
# ══════════════════════════════════════════════════════════════════════════════
h2("2.1  Overview of the Analysis Framework")
body(
    "A multi-stage machine learning pipeline was developed to predict crop maturity expressed "
    "as days after planting (DAP) from UAV-derived RGB imagery. The framework integrates three "
    "sequential components: (1) per-flight spectral feature extraction from processed UAV imagery, "
    "(2) supervised regression model training using ground-truth maturity records (MTR) collected "
    "in the field, and (3) a flight planning advisor that identifies the minimum number of UAV "
    "acquisitions required to achieve near-optimal prediction accuracy. The full pipeline was "
    "implemented in Python using scikit-learn, NumPy, pandas, and optionally XGBoost. All analyses "
    "were performed at the individual plot level, where each observation corresponds to one plot "
    "within one trial."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.2
# ══════════════════════════════════════════════════════════════════════════════
h2("2.2  Study Design and Data Sources")
body(
    "Data were collected across multiple field trials conducted at different locations and/or "
    "growing seasons. Each trial consisted of a set of experimental plots arranged according to a "
    "standard field trial design. UAV flights were conducted at regular intervals throughout the "
    "growing season, and RGB orthomosaics were processed into per-plot time-series of spectral "
    "indices. Ground-truth maturity dates were recorded by field personnel as the DAP at which each "
    "plot reached physiological maturity, and were provided in a structured Excel file (hereinafter "
    "“field Excel”) alongside plot identifiers, trial identifiers, and genotype names."
)
body(
    "To enable multi-environment analysis, each location–year combination was treated as an "
    "independent dataset defined by a root directory containing the processed UAV outputs and its "
    "own corresponding field Excel file. This structure allowed data from multiple trials and "
    "environments to be loaded independently and merged into a single combined training matrix, "
    "enabling cross-environment generalisation to be evaluated rigorously."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.3
# ══════════════════════════════════════════════════════════════════════════════
h2("2.3  UAV Image Processing and Spectral Index Derivation")
body(
    "UAV image processing — including orthomosaic generation, plot boundary delineation, and "
    "per-plot spectral index computation — was performed in a separate image analysis pipeline "
    "prior to the ML step. The outputs of this upstream pipeline used as inputs to the ML module "
    "consisted of two file types per trial:"
)
bullet(
    ": a per-plot summary table containing, for each of 23 spectral and colour-based methods, "
    "the estimated DAP at which that method’s signal crossed a predefined maturity threshold "
    "(denoted {method}_DAP). The 23 methods were: GCC, RCC, NGRDI, VARI, GLI, MGRVI, RGBVI, "
    "ExGR, IKAW, NDYI, R_over_G, TGI, WI, HMI_MASKED, HMI_MAIN, MPI, desicc_frac, green_cover, "
    "Lab_a, Lab_b, Lab_Chroma, Lab_HueAngle, and hist_ratio.",
    bold_prefix="SUMMARY.xlsx"
)
bullet(
    ": per-plot time-series files containing flight-by-flight observations of spectral indices "
    "(DAP, ExGR, GCC, RCC, HMI_MASKED, and others) used to derive temporal trajectory features.",
    bold_prefix="TimeSeries/ts_{PlotID}.csv"
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.4
# ══════════════════════════════════════════════════════════════════════════════
h2("2.4  Feature Engineering")

h3("2.4.1  Per-Method Maturity Estimates (23 features)")
body(
    "The primary feature set consisted of the 23 method-specific maturity DAP estimates read "
    "directly from SUMMARY.xlsx. Each value represented the DAP at which the corresponding "
    "spectral index crossed a method-specific threshold indicative of crop senescence or maturity. "
    "Missing estimates (e.g., when a threshold was never crossed within the season) were retained "
    "as missing values and imputed downstream during preprocessing."
)

h3("2.4.2  Consensus Statistics (4 features)")
body(
    "Three cross-method consensus statistics were computed from the 23 method-level estimates "
    "for each plot: the mean (consensus_mean), median (consensus_median), and standard deviation "
    "(consensus_std) across all methods with finite estimates. The number of methods that produced "
    "a valid estimate was also recorded (n_detected). These features captured the central tendency "
    "and spread of the multi-method ensemble, serving as robust aggregate predictors even when "
    "individual methods failed."
)

h3("2.4.3  Time-Series Derived Features (8 features)")
body(
    "Additional temporal trajectory features were computed from the per-plot time-series files "
    "using PCHIP (Piecewise Cubic Hermite Interpolating Polynomial) interpolation and ordinary "
    "least-squares linear regression:"
)
bullet(
    ": the DAP at which the Excess Green (ExGR) index reached its seasonal maximum, estimated "
    "by interpolating the ExGR time series to a 500-point dense grid and locating the global maximum.",
    bold_prefix="ExG peak DAP (exg_peak_dap)"
)
bullet(
    ": the linear rate of ExGR decline from peak to the final observation, computed as "
    "(ExGR_last − ExGR_peak) / (DAP_last − DAP_peak). A more negative value indicates "
    "faster canopy senescence.",
    bold_prefix="ExG decline slope (exg_slope)"
)
bullet(
    ": the ordinary least-squares slope of the Green Chromatic Coordinate (GCC) over the "
    "entire recorded season. A negative slope indicates progressive canopy browning.",
    bold_prefix="GCC drop rate (gcc_drop_rate)"
)
bullet(
    ": the ordinary least-squares slope of the Red Chromatic Coordinate (RCC) over the season. "
    "An increasing RCC is associated with ripening.",
    bold_prefix="RCC rise rate (rcc_rise_rate)"
)
bullet(
    ": the DAP at which HMI_MASKED first exceeded 0.80, estimated by PCHIP interpolation "
    "with 500 evaluation points.",
    bold_prefix="HMI crossing DAP (hmi_crossing_dap)"
)
bullet(
    ": the total number of valid flight acquisitions for the plot, the DAP of the first and "
    "last acquisition, and the total temporal span in days. These features captured data quality "
    "and seasonal coverage.",
    bold_prefix="Flight metadata (n_flights, first_flight_dap, last_flight_dap, flight_span_days)"
)
body(
    "The complete feature set thus comprised 38 continuous predictors plus one categorical "
    "variable (genotype name), for a total of up to 39 input features per observation."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.5
# ══════════════════════════════════════════════════════════════════════════════
h2("2.5  Multi-Environment Data Consolidation")
body(
    "When data from multiple locations or growing seasons were included, each dataset was "
    "processed independently through the feature extraction and field-data joining steps. Trial "
    "names were prefixed with a dataset label (e.g., Location2021/TrialName) to preserve the "
    "origin of each observation. Before concatenation, column structures across datasets were "
    "aligned by taking the union of all column names and filling missing columns with NaN. The "
    "combined DataFrame was saved as merged_trials_locations.xlsx for audit purposes before being "
    "passed to the training pipeline."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.6
# ══════════════════════════════════════════════════════════════════════════════
h2("2.6  Field Data Integration")
body(
    "Ground-truth MTR values were joined to the extracted feature matrix via an inner join "
    "between the ML database (UAV features) and the field Excel (ground-truth records). PlotID "
    "matching used a normalisation procedure that converted all PlotID strings to canonical "
    "integer form (e.g., “101.0” → “101”) to accommodate format "
    "inconsistencies between GIS-derived identifiers and field book entries. Trial matching used "
    "the first numeric subsequence extracted from each trial identifier (e.g., "
    "“SEVREC2025” → “2025”), enabling fuzzy alignment across naming "
    "conventions. Three cascade join strategies were attempted in order: (1) normalised PlotID "
    "plus trial number, (2) normalised PlotID alone, and (3) normalised PlotID alone without a "
    "trial column in the field data. The first strategy to return at least one matched row was "
    "adopted. Observations with missing MTR values after joining were removed."
)
body(
    "Column names for PlotID, trial, genotype, and MTR in the field Excel were resolved through "
    "a four-level lookup: exact match, case-insensitive match, an alias table covering common "
    "synonyms (e.g., “variety”, “cultiva”, “entry_name” for genotype; "
    "“maturity_dap”, “mat” for MTR), and prefix-based fallback for "
    "plot-identifier columns. This design ensured compatibility with heterogeneous field data "
    "formats across multiple collaborating sites."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.7
# ══════════════════════════════════════════════════════════════════════════════
h2("2.7  Outlier Removal")
body(
    "Prior to model training, outliers in the target variable (field MTR) were identified and "
    "removed using the interquartile range (IQR) method. The first (Q1) and third (Q3) quartiles "
    "of the MTR distribution were computed, and the IQR was defined as Q3 − Q1. Any "
    "observation with MTR outside the interval [Q1 − 1.5 × IQR, Q3 + 1.5 × IQR] "
    "was excluded. The same row mask was applied consistently to the feature matrix, the target "
    "vector, the trial group labels, and the original combined DataFrame, ensuring no misalignment "
    "in any downstream step. The number of removed observations and the accepted value range were "
    "recorded in the run log."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.8
# ══════════════════════════════════════════════════════════════════════════════
h2("2.8  Feature Preprocessing")
body(
    "All continuous features were preprocessed within a scikit-learn Pipeline using a "
    "ColumnTransformer with two branches:"
)
bullet(
    ": median imputation followed by standardisation to zero mean and unit variance "
    "(StandardScaler). Median imputation was chosen over mean imputation for robustness to the "
    "skewed distributions characteristic of spectral timing estimates.",
    bold_prefix="Continuous features"
)
bullet(
    ": mode imputation followed by ordinal encoding (OrdinalEncoder), which assigned an "
    "integer rank to each genotype in alphabetical order. Unseen genotypes encountered at "
    "prediction time were assigned the value −1.",
    bold_prefix="Categorical feature (genotype name)"
)
body(
    "This preprocessing was embedded inside each model’s pipeline, ensuring that imputation "
    "and scaling parameters were fitted exclusively on training data during cross-validation, "
    "with no information leakage from test folds."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.9
# ══════════════════════════════════════════════════════════════════════════════
h2("2.9  Machine Learning Models")
body("Seven regression algorithms were trained on the full combined dataset:")

models_data = [
    ("Ridge Regression (RidgeCV)",
     "Linear L2-regularised regression with automatic regularisation strength (α) selection "
     "via internal 5-fold cross-validation over a logarithmic grid of 60 values spanning "
     "α ∈ [0.01, 10,000]."),
    ("ElasticNet (ElasticNetCV)",
     "L1+L2 penalised linear regression with simultaneous cross-validated selection of α and "
     "the L1 ratio from the grid l1_ratio ∈ {0.1, 0.3, 0.5, 0.7, 0.9, 0.95, 1.0} and α "
     "over a 40-point logarithmic grid."),
    ("Support Vector Regression (SVR)",
     "Radial basis function (RBF) kernel SVR with fixed hyperparameters C = 10, ε = 0.5, "
     "and γ = “scale” (i.e., 1 / (n_features × Var(X)))."),
    ("Random Forest (RandomForestRegressor)",
     "Ensemble of 300 decision trees with √p features considered at each split, a minimum "
     "of 2 samples per leaf, and a fixed random seed of 42."),
    ("Gradient Boosting (GradientBoostingRegressor)",
     "Sequential ensemble of 300 shallow trees (max_depth = 4) with a learning rate of 0.05 and "
     "a row-subsampling ratio of 0.8."),
    ("Extra Trees (ExtraTreesRegressor)",
     "Ensemble of 300 extremely randomised trees with identical structural hyperparameters to "
     "the Random Forest."),
    ("XGBoost (XGBRegressor; when installed)",
     "Gradient-boosted trees with 300 estimators, learning rate 0.05, max depth 4, "
     "row-subsampling 0.8, and column-subsampling 0.8."),
]
for i, (name, desc) in enumerate(models_data, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(name + ": ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

body(
    "All ensemble models used a fixed random seed (42) to ensure reproducibility. "
    "All models were trained on the outlier-filtered, merged dataset."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.10
# ══════════════════════════════════════════════════════════════════════════════
h2("2.10  Cross-Validation Strategy")
body(
    "Model generalisation was assessed using Leave-One-Trial-Out cross-validation (LOTO-CV) "
    "when three or more unique trials were present in the dataset. Under LOTO-CV, all observations "
    "belonging to one trial were withheld as the test set while all observations from the remaining "
    "trials were used for training; this was repeated until each trial had served once as the test "
    "set. This strategy directly simulates the intended deployment scenario, where a model trained "
    "on historical trials is applied to a completely new, unseen trial — a more stringent and "
    "ecologically valid evaluation than random k-fold cross-validation, which would allow "
    "observations from the same trial to appear in both training and test sets within a given fold. "
    "When fewer than three trials were available, 5-fold cross-validation was used as a fallback."
)
body(
    "For each fold, a fresh copy of the full preprocessing and regression pipeline was fitted from "
    "scratch on the training partition, ensuring that imputation statistics and scale parameters "
    "were never estimated from test data. Predictions from all test folds were assembled into a "
    "full-length prediction vector aligned to the original dataset. Overall performance was "
    "summarised by Root Mean Squared Error (RMSE), Mean Absolute Error (MAE), coefficient of "
    "determination (R²), and systematic bias (mean of predicted minus observed MTR), computed "
    "over all test-fold predictions jointly."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.11
# ══════════════════════════════════════════════════════════════════════════════
h2("2.11  Multi-Level Stacking Ensemble")
body(
    "A two-level stacking ensemble was constructed to combine the predictions of all base "
    "learners. At Level 1, out-of-fold (OOF) predictions were generated for each base model using "
    "the same LOTO-CV (or k-fold) splitting strategy applied to base-model evaluation, ensuring "
    "that the meta-learner never observed in-sample predictions during training. At Level 2, an "
    "XGBoost meta-learner (or a Random Forest fallback when XGBoost was unavailable) was trained "
    "on a combined feature matrix consisting of the original predictor variables concatenated with "
    "the Level-1 OOF predictions. Using the original features alongside the OOF predictions, "
    "rather than OOF predictions alone, provided the meta-learner with the full predictive signal "
    "rather than a compressed ensemble summary."
)
body(
    "Stacking performance was evaluated through nested cross-validation: for each outer test fold, "
    "base models were re-trained on the outer training set, inner OOF predictions were generated "
    "within the outer training set, the meta-learner was trained on the inner OOF meta-features, "
    "and the final stacking prediction for the outer test fold was produced without any exposure "
    "to test-fold data at any stage. This nested procedure yields unbiased estimates of the "
    "stacking ensemble’s generalisation error."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.12
# ══════════════════════════════════════════════════════════════════════════════
h2("2.12  Performance Metrics")
body("Model accuracy was quantified using four metrics computed over LOTO-CV test folds:")

bullet(
    " (days): the square root of the mean squared error between predicted and observed MTR, "
    "serving as the primary criterion for model selection.",
    bold_prefix="RMSE"
)
bullet(
    " (days): the mean absolute error, less sensitive to large individual prediction errors.",
    bold_prefix="MAE"
)
bullet(
    ": the coefficient of determination, measuring the proportion of variance in field MTR "
    "explained by the model.",
    bold_prefix="R²"
)
bullet(
    " (days): the mean of (predicted − observed), indicating whether the model "
    "systematically over- or under-predicts maturity.",
    bold_prefix="Bias"
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.13
# ══════════════════════════════════════════════════════════════════════════════
h2("2.13  Ridge Formula Extraction")
body(
    "To provide an interpretable, equation-form representation of the best linear model, "
    "per-genotype prediction formulas were extracted from the fitted Ridge regression pipeline. "
    "The Ridge coefficient vector was decomposed into continuous-feature terms (ranked by absolute "
    "magnitude, top 10 reported) and a genotype offset derived from the ordinal encoder coefficient "
    "and the genotype’s alphabetical rank. Each formula took the form:"
)
equation(
    "MTR ≈ (global intercept + genotype offset) + β₁ × feature₁ + β₂ × feature₂ + ... + β₁₀ × feature₁₀"
)
body(
    "This explicit formula allows practitioners to estimate maturity DAP for any genotype given a "
    "subset of spectral index estimates, without requiring a computer or the full model pipeline."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.14
# ══════════════════════════════════════════════════════════════════════════════
h2("2.14  Flight Planning Advisor")
body(
    "To support operational UAV mission planning, a flight planning simulation was conducted for "
    "each trained model, including the stacking ensemble. The simulation addressed the following "
    "question: what is the minimum number of UAV flights, and at what DAP should the last flight "
    "occur, to achieve near-optimal maturity prediction accuracy?"
)
body(
    "For each candidate flight count N (ranging from 2 to the maximum observed number of flights "
    "per plot in the training data), the time-series of every plot was truncated to its first N "
    "observations. Flight-dependent features (ExG slope, GCC/RCC rates, HMI crossing DAP, "
    "last-flight DAP, and flight span) were recomputed from the truncated series, while static "
    "features (the 23 method DAP estimates from SUMMARY.xlsx) were retained unchanged. The trained "
    "model was applied to this simulated truncated feature set, and RMSE against field MTR was "
    "computed across all plots."
)
body(
    "The value of N producing the lowest RMSE was designated the optimal flight count. The "
    "corresponding mean last-flight DAP across all plots defined the recommended stopping date. "
    "The recommended flight interval was computed as (last-flight DAP − first-flight DAP) / "
    "(N − 1). An optimal flying window was further defined as the range of last-flight DAP "
    "values corresponding to the five N values with the lowest RMSE (the top-5 window), providing "
    "operational flexibility around the exact stopping date."
)
body(
    "Results were visualised as a single flight schedule plot showing a Gaussian bell curve "
    "centred on the start-of-flying DAP, with consecutive flight strips coloured according to the "
    "cumulative probability of maturity computed from a normal distribution centred on the "
    "predicted maturity date (from light to dark blue using a sequential colormap). Vertical "
    "markers indicated the start of flying, the predicted maturity date, and the recommended "
    "end-of-flight date."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.15
# ══════════════════════════════════════════════════════════════════════════════
h2("2.15  Model Persistence and Prediction for New Trials")
body(
    "The best-performing model bundle — including the fitted model pipeline, preprocessing "
    "parameters, feature names, cross-validation results, and Ridge formula — was serialised "
    "to disk using Python’s pickle module. To predict maturity for a new, unseen trial, the "
    "feature extraction procedure (Section 2.4) was applied to the new trial’s UAV output "
    "directory, and the stored model bundle was loaded to generate per-plot maturity predictions "
    "expressed as DAP. When a sowing date was provided, predicted DAP values were converted to "
    "calendar dates. This inference step required no field MTR data, enabling fully prospective "
    "deployment of the trained model to new growing seasons."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.16
# ══════════════════════════════════════════════════════════════════════════════
h2("2.16  Software and Reproducibility")
body(
    "All analyses were conducted in Python using the following core libraries: scikit-learn "
    "(machine learning pipelines, model training, and cross-validation), XGBoost (gradient "
    "boosting and stacking meta-learner), pandas (tabular data management), NumPy (numerical "
    "computation), SciPy (PCHIP interpolation), and Matplotlib (visualisation). The full pipeline "
    "was implemented as a standalone Python module (ml_analysis.py) integrated into a Tkinter "
    "desktop application. Random seeds were fixed at 42 for all stochastic model components to "
    "ensure reproducibility. All output files — feature matrices, predictions, model bundles, "
    "diagnostic plots, and summary statistics — were written to a structured output directory "
    "with one subdirectory per trial and dataset."
)

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
